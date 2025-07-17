"""
Multi-Modal RAG Support with Local CLIP Integration.

This module provides comprehensive multi-modal support for the RAG pipeline,
including image, document, and code understanding using local models only.
All processing is performed locally with zero external API calls.
"""

import asyncio
from typing import List, Dict, Optional, Tuple, Union, Any, BinaryIO
from datetime import datetime
from dataclasses import dataclass, field
from enum import Enum
import numpy as np
from pathlib import Path
import mimetypes
import base64
from io import BytesIO

# Document processing imports
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import cv2

# Code processing imports
from pygments import highlight
from pygments.lexers import get_lexer_by_name, guess_lexer
from pygments.formatters import TextFormatter
from pygments.util import ClassNotFound

# CLIP for image embeddings
try:
    import clip
    import torch
    CLIP_AVAILABLE = True
except ImportError:
    CLIP_AVAILABLE = False

import structlog
from pydantic import BaseModel, Field, ConfigDict
import json

from ...models.memory import Memory
from ..embeddings.embedder import Embedder
from ..cache.redis_cache import RedisCache
from ..utils.config import settings

logger = structlog.get_logger(__name__)


class ModalityType(str, Enum):
    """Types of modalities supported."""
    TEXT = "text"
    IMAGE = "image"
    DOCUMENT = "document"
    CODE = "code"
    AUDIO = "audio"  # For future expansion
    VIDEO = "video"  # For future expansion


class DocumentType(str, Enum):
    """Types of documents supported."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    HTML = "html"
    RTF = "rtf"


class CodeLanguage(str, Enum):
    """Programming languages supported."""
    PYTHON = "python"
    JAVASCRIPT = "javascript"
    TYPESCRIPT = "typescript"
    JAVA = "java"
    CPP = "cpp"
    CSHARP = "csharp"
    GO = "go"
    RUST = "rust"
    RUBY = "ruby"
    PHP = "php"
    SQL = "sql"
    JSON = "json"
    YAML = "yaml"
    XML = "xml"
    HTML = "html"
    CSS = "css"


@dataclass
class MultiModalContent:
    """Container for multi-modal content."""
    modality: ModalityType
    content: Union[str, bytes, np.ndarray]
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[np.ndarray] = None
    confidence: float = 1.0


class MultiModalSearchResult(BaseModel):
    """Result from multi-modal search."""
    model_config = ConfigDict(arbitrary_types_allowed=True)
    
    memory_id: str = Field(description="Memory ID")
    modality: ModalityType = Field(description="Content modality")
    similarity_score: float = Field(ge=0.0, le=1.0, description="Similarity score")
    content_preview: str = Field(description="Preview of the content")
    metadata: Dict[str, Any] = Field(default_factory=dict, description="Additional metadata")


class MultiModalProcessor:
    """
    Advanced multi-modal content processor for RAG pipeline.
    
    Features:
    - Local CLIP integration for image understanding
    - Document parsing (PDF, DOCX, TXT, MD, HTML)
    - Code syntax analysis with pygments
    - Unified embedding generation for all modalities
    - Multi-modal similarity search
    - Content extraction and preview generation
    """
    
    def __init__(
        self,
        embedder: Embedder,
        cache: Optional[RedisCache] = None,
        clip_model_name: str = "ViT-B/32",
        max_image_size: Tuple[int, int] = (224, 224)
    ):
        """
        Initialize the multi-modal processor.
        
        Args:
            embedder: Text embedder for non-image content
            cache: Optional Redis cache for performance
            clip_model_name: CLIP model to use for image embeddings
            max_image_size: Maximum image size for processing
        """
        self.embedder = embedder
        self.cache = cache
        self.max_image_size = max_image_size
        
        # Initialize CLIP model if available
        self.clip_model = None
        self.clip_preprocess = None
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
        if CLIP_AVAILABLE:
            self._initialize_clip(clip_model_name)
        
        # Initialize code lexers cache
        self.lexers_cache = {}
        
        logger.info(
            "Initialized multi-modal processor",
            clip_available=CLIP_AVAILABLE,
            device=self.device,
            max_image_size=max_image_size
        )
    
    def _initialize_clip(self, model_name: str):
        """Initialize CLIP model for image processing."""
        try:
            self.clip_model, self.clip_preprocess = clip.load(
                model_name, 
                device=self.device
            )
            self.clip_model.eval()
            logger.info(f"Loaded CLIP model: {model_name}")
        except Exception as e:
            logger.warning(f"Failed to load CLIP model: {e}")
            self.clip_model = None
    
    async def process_content(
        self, 
        content: Union[str, bytes, Path], 
        content_type: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> MultiModalContent:
        """
        Process content of any modality.
        
        Args:
            content: Content to process (text, bytes, or file path)
            content_type: MIME type or explicit type
            metadata: Additional metadata
            
        Returns:
            Processed multi-modal content with embeddings
        """
        if metadata is None:
            metadata = {}
        
        # Determine modality and content type
        modality = await self._detect_modality(content, content_type)
        
        # Process based on modality
        if modality == ModalityType.TEXT:
            processed = await self._process_text(content, metadata)
        elif modality == ModalityType.IMAGE:
            processed = await self._process_image(content, metadata)
        elif modality == ModalityType.DOCUMENT:
            processed = await self._process_document(content, content_type, metadata)
        elif modality == ModalityType.CODE:
            processed = await self._process_code(content, metadata)
        else:
            # Fallback to text processing
            processed = await self._process_text(str(content), metadata)
        
        return processed
    
    async def _detect_modality(
        self, 
        content: Union[str, bytes, Path], 
        content_type: Optional[str] = None
    ) -> ModalityType:
        """Detect the modality of the content."""
        
        # If explicit content type provided
        if content_type:
            if content_type.startswith('image/'):
                return ModalityType.IMAGE
            elif content_type.startswith('text/'):
                return ModalityType.TEXT
            elif content_type in ['application/pdf', 'application/msword']:
                return ModalityType.DOCUMENT
        
        # If it's a file path
        if isinstance(content, (str, Path)) and Path(content).exists():
            path = Path(content)
            mime_type, _ = mimetypes.guess_type(str(path))
            
            if mime_type:
                if mime_type.startswith('image/'):
                    return ModalityType.IMAGE
                elif mime_type == 'application/pdf':
                    return ModalityType.DOCUMENT
                elif path.suffix.lower() in ['.docx', '.doc']:
                    return ModalityType.DOCUMENT
                elif path.suffix.lower() in ['.py', '.js', '.java', '.cpp', '.c', '.h']:
                    return ModalityType.CODE
        
        # If it's bytes, try to detect image
        if isinstance(content, bytes):
            try:
                Image.open(BytesIO(content))
                return ModalityType.IMAGE
            except:
                pass
        
        # If it's text, check for code patterns
        if isinstance(content, str):
            if self._looks_like_code(content):
                return ModalityType.CODE
        
        # Default to text
        return ModalityType.TEXT
    
    def _looks_like_code(self, text: str) -> bool:
        """Heuristic to detect if text looks like code."""
        code_indicators = [
            'def ', 'function ', 'class ', 'import ', 'from ',
            '#!/', '<?php', '<?xml', '<html', '<!DOCTYPE',
            'SELECT ', 'INSERT ', 'UPDATE ', 'DELETE ',
            '{', '}', ';', '//', '/*', '*/', '#include'
        ]
        
        text_lower = text.lower()
        indicator_count = sum(1 for indicator in code_indicators if indicator in text_lower)
        
        # If multiple indicators or specific patterns, likely code
        return indicator_count >= 2 or any(
            pattern in text for pattern in ['```', '    def ', '    function ']
        )
    
    async def _process_text(
        self, 
        content: str, 
        metadata: Dict[str, Any]
    ) -> MultiModalContent:
        """Process text content."""
        # Generate text embedding
        embedding = await self.embedder.embed(content)
        
        return MultiModalContent(
            modality=ModalityType.TEXT,
            content=content,
            metadata={
                **metadata,
                'length': len(content),
                'word_count': len(content.split()),
                'processed_at': datetime.utcnow().isoformat()
            },
            embedding=embedding,
            confidence=1.0
        )
    
    async def _process_image(
        self, 
        content: Union[str, bytes, Path], 
        metadata: Dict[str, Any]
    ) -> MultiModalContent:
        """Process image content using CLIP."""
        # Load image
        if isinstance(content, (str, Path)):
            image = Image.open(content)
        elif isinstance(content, bytes):
            image = Image.open(BytesIO(content))
        else:
            raise ValueError("Invalid image content type")
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Resize if too large
        if image.size[0] > self.max_image_size[0] or image.size[1] > self.max_image_size[1]:
            image.thumbnail(self.max_image_size, Image.Resampling.LANCZOS)
        
        # Generate image embedding
        embedding = None
        confidence = 0.5  # Default confidence for images without CLIP
        
        if self.clip_model and CLIP_AVAILABLE:
            try:
                # Preprocess image for CLIP
                image_tensor = self.clip_preprocess(image).unsqueeze(0).to(self.device)
                
                # Generate embedding
                with torch.no_grad():
                    embedding = self.clip_model.encode_image(image_tensor)
                    embedding = embedding.cpu().numpy().flatten()
                    confidence = 1.0
                    
            except Exception as e:
                logger.warning(f"CLIP processing failed: {e}")
        
        # Fallback: convert image to text description for text embedding
        if embedding is None:
            image_description = f"Image with dimensions {image.size}, format: {image.format}"
            embedding = await self.embedder.embed(image_description)
        
        # Convert image to bytes for storage
        img_bytes = BytesIO()
        image.save(img_bytes, format='PNG')
        img_bytes = img_bytes.getvalue()
        
        return MultiModalContent(
            modality=ModalityType.IMAGE,
            content=img_bytes,
            metadata={
                **metadata,
                'width': image.size[0],
                'height': image.size[1],
                'format': image.format,
                'mode': image.mode,
                'size_bytes': len(img_bytes),
                'processed_at': datetime.utcnow().isoformat()
            },
            embedding=embedding,
            confidence=confidence
        )
    
    async def _process_document(
        self, 
        content: Union[str, bytes, Path], 
        content_type: Optional[str],
        metadata: Dict[str, Any]
    ) -> MultiModalContent:
        """Process document content (PDF, DOCX, etc.)."""
        text_content = ""
        doc_metadata = {}
        
        try:
            if isinstance(content, (str, Path)):
                file_path = Path(content)
                
                if file_path.suffix.lower() == '.pdf':
                    text_content, doc_metadata = await self._extract_pdf_content(file_path)
                elif file_path.suffix.lower() in ['.docx', '.doc']:
                    text_content, doc_metadata = await self._extract_docx_content(file_path)
                else:
                    # Fallback: read as text
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text_content = f.read()
            
            elif isinstance(content, bytes):
                # Try to detect document type and extract
                if content_type == 'application/pdf':
                    text_content, doc_metadata = await self._extract_pdf_from_bytes(content)
                else:
                    # Fallback: decode as text
                    text_content = content.decode('utf-8', errors='ignore')
            
            else:
                text_content = str(content)
            
        except Exception as e:
            logger.warning(f"Document processing failed: {e}")
            text_content = f"Failed to process document: {str(e)}"
        
        # Generate text embedding for document content
        embedding = await self.embedder.embed(text_content)
        
        return MultiModalContent(
            modality=ModalityType.DOCUMENT,
            content=text_content,
            metadata={
                **metadata,
                **doc_metadata,
                'length': len(text_content),
                'word_count': len(text_content.split()),
                'processed_at': datetime.utcnow().isoformat()
            },
            embedding=embedding,
            confidence=0.9 if text_content and len(text_content) > 10 else 0.3
        )
    
    async def _extract_pdf_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from PDF file."""
        text_parts = []
        metadata = {}
        
        try:
            doc = fitz.open(str(file_path))
            
            metadata.update({
                'page_count': doc.page_count,
                'metadata': doc.metadata,
                'file_size': file_path.stat().st_size
            })
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"Page {page_num + 1}:\n{text}")
            
            doc.close()
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
        
        return '\n\n'.join(text_parts), metadata
    
    async def _extract_pdf_from_bytes(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from PDF bytes."""
        text_parts = []
        metadata = {}
        
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            
            metadata.update({
                'page_count': doc.page_count,
                'metadata': doc.metadata,
                'size_bytes': len(content)
            })
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"Page {page_num + 1}:\n{text}")
            
            doc.close()
            
        except Exception as e:
            logger.error(f"PDF extraction from bytes failed: {e}")
            raise
        
        return '\n\n'.join(text_parts), metadata
    
    async def _extract_docx_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text content from DOCX file."""
        text_parts = []
        metadata = {}
        
        try:
            doc = Document(str(file_path))
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(' | '.join(row_text))
                if table_text:
                    text_parts.append('\n'.join(table_text))
            
            metadata.update({
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'file_size': file_path.stat().st_size
            })
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
        
        return '\n\n'.join(text_parts), metadata
    
    async def _process_code(
        self, 
        content: Union[str, Path], 
        metadata: Dict[str, Any]
    ) -> MultiModalContent:
        """Process code content with syntax highlighting and analysis."""
        
        # Read code content
        if isinstance(content, Path):
            with open(content, 'r', encoding='utf-8', errors='ignore') as f:
                code_text = f.read()
            file_extension = content.suffix.lower()
        else:
            code_text = str(content)
            file_extension = metadata.get('file_extension', '.txt')
        
        # Detect programming language
        language = await self._detect_code_language(code_text, file_extension)
        
        # Extract code features
        code_features = await self._extract_code_features(code_text, language)
        
        # Create enhanced text representation for embedding
        enhanced_text = await self._create_code_embedding_text(
            code_text, 
            language, 
            code_features
        )
        
        # Generate embedding
        embedding = await self.embedder.embed(enhanced_text)
        
        return MultiModalContent(
            modality=ModalityType.CODE,
            content=code_text,
            metadata={
                **metadata,
                'language': language,
                'line_count': len(code_text.split('\n')),
                'char_count': len(code_text),
                **code_features,
                'processed_at': datetime.utcnow().isoformat()
            },
            embedding=embedding,
            confidence=0.95
        )
    
    async def _detect_code_language(
        self, 
        code_text: str, 
        file_extension: str
    ) -> str:
        """Detect programming language from code and file extension."""
        
        # Extension-based detection
        extension_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.h': 'c',
            '.cs': 'csharp',
            '.go': 'go',
            '.rs': 'rust',
            '.rb': 'ruby',
            '.php': 'php',
            '.sql': 'sql',
            '.json': 'json',
            '.yaml': 'yaml',
            '.yml': 'yaml',
            '.xml': 'xml',
            '.html': 'html',
            '.css': 'css'
        }
        
        if file_extension in extension_map:
            return extension_map[file_extension]
        
        # Content-based detection using pygments
        try:
            lexer = guess_lexer(code_text)
            return lexer.name.lower()
        except ClassNotFound:
            return 'text'
    
    async def _extract_code_features(
        self, 
        code_text: str, 
        language: str
    ) -> Dict[str, Any]:
        """Extract features from code for better understanding."""
        features = {}
        
        lines = code_text.split('\n')
        
        # Basic metrics
        features.update({
            'total_lines': len(lines),
            'non_empty_lines': len([line for line in lines if line.strip()]),
            'comment_lines': len([line for line in lines if line.strip().startswith(('#', '//', '/*'))]),
        })
        
        # Language-specific feature extraction
        if language == 'python':
            features.update(self._extract_python_features(code_text))
        elif language in ['javascript', 'typescript']:
            features.update(self._extract_js_features(code_text))
        elif language == 'java':
            features.update(self._extract_java_features(code_text))
        
        return features
    
    def _extract_python_features(self, code_text: str) -> Dict[str, Any]:
        """Extract Python-specific features."""
        import re
        
        features = {}
        
        # Function definitions
        functions = re.findall(r'def\s+(\w+)', code_text)
        features['functions'] = functions
        features['function_count'] = len(functions)
        
        # Class definitions
        classes = re.findall(r'class\s+(\w+)', code_text)
        features['classes'] = classes
        features['class_count'] = len(classes)
        
        # Imports
        imports = re.findall(r'(?:from\s+\w+\s+)?import\s+([^\n]+)', code_text)
        features['imports'] = imports
        features['import_count'] = len(imports)
        
        return features
    
    def _extract_js_features(self, code_text: str) -> Dict[str, Any]:
        """Extract JavaScript/TypeScript-specific features."""
        import re
        
        features = {}
        
        # Function definitions
        functions = re.findall(r'function\s+(\w+)|(\w+)\s*=\s*(?:async\s+)?(?:\([^)]*\)\s*=>|\([^)]*\)\s*{)', code_text)
        func_names = [f[0] or f[1] for f in functions if f[0] or f[1]]
        features['functions'] = func_names
        features['function_count'] = len(func_names)
        
        # Class definitions
        classes = re.findall(r'class\s+(\w+)', code_text)
        features['classes'] = classes
        features['class_count'] = len(classes)
        
        return features
    
    def _extract_java_features(self, code_text: str) -> Dict[str, Any]:
        """Extract Java-specific features."""
        import re
        
        features = {}
        
        # Method definitions
        methods = re.findall(r'(?:public|private|protected|static|\s)+[\w<>\[\]]+\s+(\w+)\s*\([^)]*\)\s*{', code_text)
        features['methods'] = methods
        features['method_count'] = len(methods)
        
        # Class definitions
        classes = re.findall(r'(?:public|abstract|final|\s)*class\s+(\w+)', code_text)
        features['classes'] = classes
        features['class_count'] = len(classes)
        
        return features
    
    async def _create_code_embedding_text(
        self, 
        code_text: str, 
        language: str, 
        features: Dict[str, Any]
    ) -> str:
        """Create enhanced text representation for code embedding."""
        
        # Start with language and basic description
        parts = [f"Programming language: {language}"]
        
        # Add feature summary
        if features.get('function_count', 0) > 0:
            parts.append(f"Contains {features['function_count']} functions")
        
        if features.get('class_count', 0) > 0:
            parts.append(f"Contains {features['class_count']} classes")
        
        # Add code structure summary
        parts.append(f"Code structure: {features.get('total_lines', 0)} lines")
        
        # Add function/class names for searchability
        if features.get('functions'):
            parts.append(f"Functions: {', '.join(features['functions'][:5])}")
        
        if features.get('classes'):
            parts.append(f"Classes: {', '.join(features['classes'][:5])}")
        
        # Add a portion of the actual code for content matching
        code_preview = code_text[:500] + ('...' if len(code_text) > 500 else '')
        parts.append(f"Code content: {code_preview}")
        
        return '\n'.join(parts)
    
    async def search_multimodal(
        self,
        query: str,
        memories: List[Memory],
        modalities: Optional[List[ModalityType]] = None,
        top_k: int = 10
    ) -> List[MultiModalSearchResult]:
        """
        Search across multiple modalities.
        
        Args:
            query: Search query
            memories: List of memories to search
            modalities: List of modalities to include in search
            top_k: Number of top results to return
            
        Returns:
            List of search results sorted by relevance
        """
        if modalities is None:
            modalities = list(ModalityType)
        
        # Generate query embedding
        query_embedding = await self.embedder.embed(query)
        
        results = []
        
        for memory in memories:
            # Skip if memory doesn't match desired modalities
            memory_modality = memory.metadata.get('modality', ModalityType.TEXT)
            if memory_modality not in modalities:
                continue
            
            # Calculate similarity
            if hasattr(memory, 'embedding') and memory.embedding is not None:
                similarity = self._calculate_similarity(query_embedding, memory.embedding)
            else:
                # Fallback: use content similarity
                similarity = await self._calculate_content_similarity(query, memory.content)
            
            # Create content preview
            preview = self._create_content_preview(memory.content, memory_modality)
            
            result = MultiModalSearchResult(
                memory_id=memory.id,
                modality=memory_modality,
                similarity_score=similarity,
                content_preview=preview,
                metadata=memory.metadata
            )
            
            results.append(result)
        
        # Sort by similarity and return top-k
        results.sort(key=lambda x: x.similarity_score, reverse=True)
        return results[:top_k]
    
    def _calculate_similarity(
        self, 
        embedding1: np.ndarray, 
        embedding2: np.ndarray
    ) -> float:
        """Calculate cosine similarity between embeddings."""
        from sklearn.metrics.pairwise import cosine_similarity
        
        return float(cosine_similarity([embedding1], [embedding2])[0][0])
    
    async def _calculate_content_similarity(self, query: str, content: str) -> float:
        """Fallback content similarity calculation."""
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        
        try:
            vectorizer = TfidfVectorizer()
            vectors = vectorizer.fit_transform([query, str(content)])
            similarity = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
            return float(similarity)
        except:
            return 0.0
    
    def _create_content_preview(
        self, 
        content: Union[str, bytes], 
        modality: ModalityType
    ) -> str:
        """Create a preview of the content for display."""
        
        if modality == ModalityType.IMAGE:
            return "[Image content]"
        
        elif modality == ModalityType.DOCUMENT:
            text_content = str(content)
            preview = text_content[:200]
            return preview + ('...' if len(text_content) > 200 else '')
        
        elif modality == ModalityType.CODE:
            code_content = str(content)
            lines = code_content.split('\n')
            preview_lines = lines[:5]
            preview = '\n'.join(preview_lines)
            if len(lines) > 5:
                preview += f'\n... ({len(lines) - 5} more lines)'
            return preview
        
        else:  # TEXT
            text_content = str(content)
            preview = text_content[:200]
            return preview + ('...' if len(text_content) > 200 else '')
    
    async def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get list of supported formats by modality."""
        return {
            ModalityType.TEXT.value: ['txt', 'md', 'rtf'],
            ModalityType.IMAGE.value: ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff'],
            ModalityType.DOCUMENT.value: ['pdf', 'docx', 'doc', 'html'],
            ModalityType.CODE.value: [
                'py', 'js', 'ts', 'java', 'cpp', 'c', 'h', 'cs', 
                'go', 'rs', 'rb', 'php', 'sql', 'json', 'yaml', 'xml', 'css'
            ]
        }